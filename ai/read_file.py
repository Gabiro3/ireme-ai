import fitz  # PyMuPDF
from docx import Document
import openpyxl
import requests
def extract_text_from_pdf(pdf_path):
    # Open the PDF file
    pdf_document = fitz.open(pdf_path)

    # Iterate through pages and extract text
    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        text = page.get_text("text")
        return f"Page {page_number + 1}:\n{text}\n"

    # Close the PDF file
    pdf_document.close()


def extract_from_word(file_path):
    # Load the Word document
    doc = Document(file_path)
    
    # Extract and print the text from the document
    for paragraph in doc.paragraphs:
        return paragraph.text
    
    # Extract and print text from tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                return cell.text
            


def summarize(text_input):
    # Define the endpoint URL
    url = "https://datacubeai.pythonanywhere.com/api/generate"
    
    # Combine the query with the text input
    contents = f"Can you summarise this letter for me and please highlight some important information from it: {text_input}"
    
    # Define the request payload
    payload = {
        "model": "gemini-pro",
        "contents": contents
    }
    headers = {
        "content-type": "application/json"
    }
    
    try:
        # Make the POST request
        response = requests.post(url, json=payload, headers=headers)
        
        # Raise an exception if the request was unsuccessful
        response.raise_for_status()
        
        # Check if the response content is empty
        if not response.content:
            print("Empty response content")
            return 'No summary available'
        
        # Parse the JSON response
        response_data = response.json()
        
        # Assuming the response has a key 'summary' which holds the AI generated text
        return response_data.get('response', 'No summary available')
    
    except requests.exceptions.RequestException as e:
        # Print any error that occurs
        print(f"An error occurred: {e}")
        return 'Error in generating summary'
    except ValueError as e:
        # Handle JSON decoding error
        print(f"JSON decoding error: {e}")
        print(f"Response content: {response.content}")
        return 'Error in generating summary'

def tag(text_input):
    # Define the endpoint URL
    url = "https://datacubeai.pythonanywhere.com/api/generate"
    
    # Combine the query with the text input
    contents = f"Hi, give this text  some tags like #recommendation, #student etc. that highlight the content inside it: {text_input}"
    
    # Define the request payload
    payload = {
        "model": "gemini-pro",
        "contents": contents
    }
    headers = {
        "content-type": "application/json"
    }
    
    try:
        # Make the POST request
        response = requests.post(url, json=payload, headers=headers)
        
        # Raise an exception if the request was unsuccessful
        response.raise_for_status()
        
        # Check if the response content is empty
        if not response.content:
            print("Empty response content")
            return 'No summary available'
        
        # Parse the JSON response
        response_data = response.json()
        
        # Assuming the response has a key 'summary' which holds the AI generated text
        return response_data.get('response', 'No summary available')
    
    except requests.exceptions.RequestException as e:
        # Print any error that occurs
        print(f"An error occurred: {e}")
        return 'Error in generating summary'
    except ValueError as e:
        # Handle JSON decoding error
        print(f"JSON decoding error: {e}")
        print(f"Response content: {response.content}")
        return 'Error in generating summary'



def extract_excel_data(file_path, sheet_name=None):
    """
    Extracts data from an Excel file and returns it as a formatted text.
    
    Args:
    - file_path: str, path to the Excel file.
    - sheet_name: str, name of the sheet to read (optional). If None, reads the active sheet.
    
    Returns:
    - str, formatted text of the sheet's data.
    """
    # Load the workbook
    workbook = openpyxl.load_workbook(file_path)

    # Get the sheet
    if sheet_name:
        sheet = workbook[sheet_name]
    else:
        sheet = workbook.active

    # Extract data
    data = []
    for row in sheet.iter_rows(values_only=True):
        data.append(row)

    # Format data as text
    formatted_data = ""
    for row in data:
        formatted_row = "\t".join([str(cell) if cell is not None else "" for cell in row])
        formatted_data += formatted_row + "\n"

    return formatted_data

